
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    doc = Document()
    
    # Define basic styles if needed, or rely on defaults
    # Setting a base font style for normal text
    # style = doc.styles['Normal']
    # font = style.font
    # font.name = 'Times New Roman'
    # font.size = Pt(13)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip()
        
        # Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            if level > 9: level = 9
            doc.add_heading(text, level=level if level <= 9 else 0)
        else:
            doc.add_paragraph(line)


    doc.save(docx_path)
    print(f"Saved {docx_path}")

target_dir = r"e:\Freelance\Research\D11_9_2025_Image_fixed_Detect\Project\T07FakeMediaDetect\docs\Bao cao chuan"
files = [f for f in os.listdir(target_dir) if f.lower().endswith('.md')]

if not files:
    print("No markdown files found in target directory.")
else:
    for filename in files:
        md_file = os.path.join(target_dir, filename)
        docx_file = os.path.join(target_dir, os.path.splitext(filename)[0] + ".docx")
        try:
            convert_md_to_docx(md_file, docx_file)
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Failed to convert {filename}: {e}")
