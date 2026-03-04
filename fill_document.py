import docx
from docx import Document
from docx.shared import Inches, Pt

# Mở file gốc
doc = Document('docs/Giay xac nhan danh gia thu nghiem san pham.docx')

# Tìm và điền thông tin vào các paragraphs
for para in doc.paragraphs:
    if 'Tên sản phẩm thử nghiệm:' in para.text:
        para.text = 'Tên sản phẩm thử nghiệm: Hệ thống phát hiện ảnh số đã bị chỉnh sửa dựa trên Luật Benford kết hợp thuật toán học máy có giám sát (T07FakeMediaDetect)'
    if 'Sinh viên thực hiện:' in para.text:
        para.text = 'Sinh viên thực hiện: Trần Quốc Bảo - Lớp B1D11'

# Tìm bảng chức năng (bảng thứ 2)
if len(doc.tables) >= 2:
    func_table = doc.tables[1]
    
    # Danh sách chức năng
    functions = [
        ('Phát hiện ảnh giả mạo bằng AI (CNN + ELA + Luật Benford)', 'Mô hình CNN + ELA + Benford, Accuracy ~94%'),
        ('Phân tích Error Level Analysis (ELA)', 'Phát hiện vùng nén khác trong ảnh JPEG'),
        ('Phát hiện cạnh (Edge Detection - Sobel)', 'Sobel operator phát hiện cạnh bất thường'),
        ('Phân tích Luminance Gradient', 'Tìm vùng không nhất quán về ánh sáng'),
        ('Phân tích nhiễu (Noise Analysis)', 'Phát hiện bất thường trong mẫu nhiễu'),
        ('Phát hiện Copy-Move (SIFT)', 'Thuật toán SIFT tìm vùng sao chép'),
        ('Tạo Binary Mask vùng giả mạo', 'Mô hình U-Net tạo mặt nạ nhị phân'),
        ('Trích xuất Metadata ảnh (EXIF)', 'Hiển thị thông tin EXIF từ ảnh'),
        ('Phát hiện video giả mạo (Frame-based CNN)', 'Phân tích từng frame phát hiện deepfake'),
        ('Trích xuất Metadata video', 'Độ phân giải, FPS, thời lượng video'),
        ('Phân tích ảnh từ PDF', 'Trích xuất và phân tích ảnh từ PDF'),
        ('Giao diện web tiếng Việt', 'Web responsive, hỗ trợ tiếng Việt'),
        ('Hệ thống quản lý file upload', 'Hỗ trợ JPG, PNG, MP4, AVI, PDF'),
        ('Hiển thị kết quả trực quan', 'Màu sắc và phần trăm tin cậy'),
        ('Phóng to ảnh kết quả', 'Click phóng to kết quả forensic'),
    ]
    
    # Xóa các hàng hiện có (trừ header)
    while len(func_table.rows) > 2:
        tr = func_table.rows[-1]._tr
        func_table._tbl.remove(tr)
    
    # Thêm các chức năng mới
    for i, (func_name, note) in enumerate(functions, 1):
        row = func_table.add_row()
        cells = row.cells
        cells[0].text = str(i)
        cells[1].text = func_name
        cells[2].text = 'X'  # Đạt
        cells[3].text = ''   # Không đạt
        cells[4].text = note

# Điền kết luận - tìm các paragraph chứa dấu chấm lửng
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('……………') and 'Kết luận' in doc.paragraphs[i-1].text if i > 0 else False:
        para.text = 'Hệ thống T07FakeMediaDetect đã hoàn thành đầy đủ 15 chức năng: Phát hiện ảnh/video giả mạo bằng AI (CNN + ELA + Luật Benford), bộ công cụ pháp y kỹ thuật số, giao diện web tiếng Việt. Độ chính xác ~94%. Đề nghị: Đạt yêu cầu.'

# Lưu file mới
output_path = 'docs/Giay_xac_nhan_DA_DIEN.docx'
doc.save(output_path)
print(f'Đã tạo file: {output_path}')
